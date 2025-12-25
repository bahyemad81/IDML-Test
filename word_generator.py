"""
Word Document Generator for Arabic Translation
Creates a .docx file with translated Arabic text in RTL format
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os


def create_word_document(idml_temp_dir, output_path):
    """
    Create a Word document from translated IDML content
    
    Args:
        idml_temp_dir: Path to extracted IDML directory
        output_path: Path where to save the Word document
        
    Returns:
        Path to created Word document
    """
    # Create new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Translated Document - Arabic', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Extract and add translated text from Stories
    stories_dir = os.path.join(idml_temp_dir, 'Stories')
    
    if os.path.exists(stories_dir):
        for story_file in sorted(os.listdir(stories_dir)):
            if story_file.endswith('.xml'):
                story_path = os.path.join(stories_dir, story_file)
                
                try:
                    # Parse the story XML
                    tree = etree.parse(story_path)
                    root = tree.getroot()
                    
                    # Find all Content elements (without namespace)
                    # IDML uses default namespace, so we need to handle it properly
                    for content in root.iter():
                        if content.tag.endswith('Content') and content.text and content.text.strip():
                            # Add paragraph with Arabic text
                            text = content.text.strip()
                            
                            # Skip very short content (likely formatting artifacts)
                            if len(text) < 2:
                                continue
                            
                            para = doc.add_paragraph(text)
                            
                            # Set paragraph to RTL
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            set_rtl_paragraph(para)
                            
                            # Set font for Arabic
                            for run in para.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                                
                                # Set RTL property for run
                                r = run._element
                                rPr = r.get_or_add_rPr()
                                rtl = OxmlElement('w:rtl')
                                rPr.append(rtl)
                
                except Exception as e:
                    print(f"Error processing story {story_file}: {e}")
                    continue
    
    # Save the document
    doc.save(output_path)
    return output_path


def set_rtl_paragraph(paragraph):
    """Set paragraph direction to RTL"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
